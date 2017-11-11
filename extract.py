import os
from statistics import mean, pvariance

from docx import Document
from docx.shared import Cm, Pt

folder = 'C:\\Users\\Kondrat\\Desktop\\Trees\\'

a = []
p = []
b = []
h = []
m = 0
d = 0


def start_extract():
    tree_files = os.listdir(folder)
    tree_files.sort(key=lambda i: int(i[:i.index('.')]))

    for treeFile in tree_files:
        if treeFile[0] != '0':
            f = open(folder + treeFile, 'r')
            a.append(float(f.readline()[:5]))
            p.append(int(f.readline()))
            b.append(int(f.readline()))
            h.append(int(f.readline()))
            f.close()

    m = mean(a)
    print('M() = ' + str(m))
    d = pvariance(a, m)
    print('D() = ' + str(d))

    create_trees_info_table()


def create_trees_info_table():
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(6)

    for i in range(len(a)):
        if (i % 16) == 0:
            table = document.add_table(rows=5, cols=0)
            document.add_paragraph()
            table.autofit = True
            table.style = 'Light Grid'
            hdr_cells = table.add_column(Cm(1)).cells
            hdr_cells[0].text = 'Number'
            hdr_cells[1].text = 'Alpha'
            hdr_cells[2].text = 'All'
            hdr_cells[3].text = 'Apex'
            hdr_cells[4].text = 'Height'
        col_cells = table.add_column(Cm(1)).cells
        col_cells[0].text = str(i + 1)
        col_cells[1].text = str(a[i])
        col_cells[2].text = str(p[i])
        col_cells[3].text = str(b[i])
        col_cells[4].text = str(h[i])
    document.add_page_break()
    document.save('demo.docx')


if __name__ == '__main__':
    start_extract()
